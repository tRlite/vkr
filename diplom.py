import docx
import docx2txt

from docx.opc.constants import RELATIONSHIP_TYPE as RT

from docx2python import docx2python

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import zipfile
import re
import aspose.words as aw

# coding: utf-8
import win32com
from win32com.client import Dispatch, DispatchEx
import os
import time
import json 
from fuzzywuzzy import fuzz
import tqdm

import pandas as pd

from collections import defaultdict

from docx.shared import Pt

from spellchecker import SpellChecker
import requests
import string
from nltk import sent_tokenize, word_tokenize

from collections import Counter
import re
import pymorphy2
from nltk import ngrams

from nltk.corpus import stopwords

from docx.text.paragraph import Paragraph
from docx.document import Document as Doc_cls
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import docx

from lxml.etree import _Element

from ruts import BasicStats, ReadabilityStats, DiversityStats

import torch
import torch.nn.functional as F

from torch import Tensor
from transformers import AutoTokenizer, AutoModel

import gc

DEVICE = 'cpu'

def average_pool(last_hidden_states: Tensor, attention_mask: Tensor) -> Tensor:
    last_hidden = last_hidden_states.masked_fill(~attention_mask[..., None].bool(), 0.0)
    return last_hidden.sum(dim=1) / attention_mask.sum(dim=1)[..., None]

def model_loading():
    tokenizer = AutoTokenizer.from_pretrained('intfloat/multilingual-e5-large')
    model = AutoModel.from_pretrained('intfloat/multilingual-e5-large').to(DEVICE)
    model.eval()
    return model, tokenizer

def semantic_scoring(model, tokenizer, header, text):
    with torch.no_grad():
        batch_dict = tokenizer([header, text], max_length=512, padding=True, truncation=True, return_tensors='pt')

    outputs = model(**batch_dict.to(DEVICE))
    embeddings = average_pool(outputs.last_hidden_state, batch_dict['attention_mask'])

    embeddings = F.normalize(embeddings, p=2, dim=1)
    score = (embeddings[0] @ embeddings[1])

    gc.collect()
    torch.cuda.empty_cache()

    return score.cpu()

def semantic_processing(doc, heading_p):
    THRS = 0.8096
    if not heading_p:
        return
    model, tokenizer = model_loading()
    for i, p in tqdm.tqdm(enumerate(doc.paragraphs), desc='Check text logic'):
        if i in heading_p:
            header = p.text
            if header.lower().strip() in {"введение", "оглавление", "заключение"}:
                continue
            inner_content = ''
            for inner_p in doc.paragraphs[i+1 : i+4]:
                inner_content += inner_p.text
            score = semantic_scoring(model, tokenizer, header, inner_content)
            print(header, inner_content, score)
            if score < THRS:
                make_comment(doc.paragraphs[i], "Название главы не соответствует ее внутреннему содержанию")

def iter_block_items(parent):
    if isinstance(parent, Doc_cls):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        print(type(child))
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)
        elif isinstance(child, _Element):
            print(''.join(child.itertext()))
            yield ''.join(child.itertext())

def check_main_elements(paragraphs):
    """Проверка наличия основных элементов в документе - [оглавление аннотация введение заключение литература]"""
    blocks = [0 for i in range(5)]

    for i, p in enumerate(paragraphs):
        p_text = p.text.strip().lower()
        if re.search(r'^\s*аннотация\s*$', p_text):
            blocks[0] = i
        if re.search(r'^\s*((оглавление)|(содержание))\s*$', p_text):
            blocks[1] = i
        if re.search(r'^\s*введение\s*$', p_text):
            blocks[2] = i
        if re.search(r'^\s*заключение\s*$', p_text):
            blocks[3] = i
        if re.search(r'^\s*список\s*.*литературы\s*$', p_text) or re.search(r'^\s*литература\s*$', p_text):
            blocks[4] = i
    
    if blocks[0] == 0:
        make_comment(paragraphs[0], 'Не найден блок аннотация')
        
    if blocks[1] == 0:
        make_comment(paragraphs[0], 'Не найден блок оглавление')

    elif blocks[1] < blocks[0]:
        make_comment(paragraphs[blocks[1]], 'Блок оглавления должен идти после блока аннотации') 
    
    if blocks[2] == 0:
        make_comment(paragraphs[0], 'Не найден блок введение')
    
    elif blocks[2] < blocks[1]:
        make_comment(paragraphs[blocks[2]], 'Блок введение должен идти после блока оглавления') 
    
    if blocks[3] == 0:
        make_comment(paragraphs[0], 'Не найден блок заключение')

    elif blocks[3] < blocks[2]:
        make_comment(paragraphs[blocks[1]], 'Блок заключение не должен идти до блока введение') 
    
    if blocks[4] == 0:
        make_comment(paragraphs[0], 'Не найден блок литературы')

    elif blocks[4] < blocks[3]:
        make_comment(paragraphs[blocks[1]], 'Блок литературы должен идти после блока заключение') 

    return blocks

def split_by_pages(path):
    """Разбиение документа по страницам с извлечением содержимого"""
    splitted_text = []
    word = Dispatch('word.application') # Open Word application
    try:
        word.visible = 0 # Run in the background, not displayed
        word.displayalerts = 0 # No warning
        PATH = os.path.abspath('C:/Users/Тагир/OneDrive/Рабочий стол/diplom/' + path)
        try:
            doc = word.Documents.Open(FileName=PATH)
        except Exception as e:
            print("BAD CONNECTION", e)
            return

        page_count = doc.ActiveWindow.Panes(1).Pages.Count
        for page_n in tqdm.tqdm(range(1, page_count + 1), desc='Split by pages'):
            rect = doc.ActiveWindow.Panes(1).Pages(page_n).Rectangles
            page_text = []
            for j in range(rect.Count):
                try:
                    page_text.append(rect.Item(j+1).Range.Text.replace("\n", " ").replace("\r", " ").replace("\t", " "))
                except Exception as e:
                    pass
            page_text = re.sub('^\d+\s', '', ''.join(page_text).strip()).strip()
            splitted_text.append(page_text)

    except Exception as e:
        print(e)
    finally:
        doc.Close()
        word.Quit
    return splitted_text

def make_comment(paragraph, comment=''):
    """
    Добавление комментариев к соответствующим элементам
    """ 
    paragraph.add_comment(comment, author='bot')

def extract_text_from_docx(docx_filename):
    """Извлечение текста из документа"""
    doc = Document(docx_filename)
    full_text = []
    for paragraph in doc.paragraphs:
        for link in paragraph._element.xpath(".//w:hyperlink"):
            inner_run = link.xpath("w:r", namespaces=link.nsmap)[0]
            print(inner_run.text)
            print(paragraph.text)
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

"""Margin of pages"""
def check_section(doc):
    """Проверка правильности оформления страницы по полям"""
    flag = False
    for section in doc.sections[1:]:
        if section.left_margin != Pt(85.05):  # 3 см
            flag = True
        if section.right_margin != Pt(42.55):  # 1,5 см
            flag = True
        if section.top_margin != Pt(56.7):  # 2 см
            flag = True
        if section.bottom_margin != Pt(56.7):  # 2 см
            flag = True
        if flag:
            section.footer.add_paragraph("""Проверьте оформление страницы. В настройках, в параметрах страницы установите следующие значения: левое поле - 3cm, правое поле: 1.5cm, верхнее поле - 2cm, нижнее поле - 2cm""")
            #doc.paragraphs[par_with_message].add_footnote("""Проверьте оформление страницы. В настройках, в параметрах страницы установите следующие значения: левое поле - 3cm, правое поле: 1.5cm, верхнее поле - 2cm, нижнее поле - 2cm""")    
        flag = False

"""Alignent and font"""
def check_font(paragraph):
    """Проверка шрифта"""
    for run in paragraph.runs:
        font = run.font
        if (run.text.replace("\n", " ").replace("\r", " ").replace("\t", " ").strip() and
           ((font.name and font.name != "Times New Roman") or (font.size and font.size != 177800))):  # 177800 единиц соответствуют 14 кеглю
            make_comment(paragraph, "Ошибка: Некорректный шрифт")
            return

def check_aligment(paragraph):
    """Проверка отступов и выравнивания в написанном тексте"""
    if paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        line_spacing = paragraph.paragraph_format.line_spacing
        if line_spacing != 1.5 and paragraph.text.strip():
            make_comment(paragraph, "Ошибка: Некорректное расстояние между строками")
    elif paragraph.text.strip():
        make_comment(paragraph, "Выравнивание текста должно быть по ширине")

"""Check title page"""

def check_title_page(doc, page_start_run):
    """Проверка оформления титульного листа"""
    if page_start_run[0] is None:
        return
    url = 'https://ru.wikipedia.org/wiki/%D0%A1%D0%BF%D0%B8%D1%81%D0%BE%D0%BA_%D0%B3%D0%BE%D1%80%D0%BE%D0%B4%D0%BE%D0%B2_%D0%A0%D0%BE%D1%81%D1%81%D0%B8%D0%B8'
    df = pd.read_html(url)[0]
    cities = list(df['Город'])
    cities.extend([ 'Севастополь', 'Евпатория', 'Саки', 'Черноморское', 'Джанкой', 'Красноперекопск',
                    'Армянск', 'Симферополь', 'Белогорск', 'Ялта', 'Алупка', 'Алушта', 'Форос', 
                    'Гурзуф', 'Партенит', 'Феодосия', 'Судак', 'Старый Крым', 'Приморский', 'Коктебель',
                    'Керчь', 'Инкерман', 'Балаклава', 'Бахчисарай', 'Донецк', 'Донецк', 'Луганск', 'Мариуполь',
                    'Бердянск', 'Мелитополь'
                  ])

    flag_minest = False

    flag_un = False
    un_err_point = 0

    flag_facul = False
    facul_err_point = 0

    flag_kaf = False
    kaf_err_point = 0

    flag_vkr = False
    vkr_err_point = 0

    flag_author = False
    author_err_point = 0

    flag_ruk = False
    ruk_err_point = 0

    flag_city = False
    flag_year = False

    next_page = 1
    while next_page < len(page_start_run) and page_start_run[next_page] is None:
        next_page += 1

    for i, p in enumerate(tqdm.tqdm(doc.paragraphs[:page_start_run[next_page][1]], desc='Process title page')):
        if 'министерство' in p.text.lower() and not flag_minest:
            flag_minest = True
            if any(e.islower() for e in p.text):
                make_comment(p, "Информация о министерстве образования должна быть введена капслоком")
            un_err_point = i + 1
        
        if 'университет' in p.text.lower() or 'институт' in p.text.lower() and not flag_un:
            flag_un = True
            if any(e.islower() for e in p.text):
                make_comment(p, "Информация об учебном заведении должна быть введена капслоком")
            facul_err_point = i + 2

        if 'факультет' in p.text.lower():
            flag_facul = True
            kaf_err_point = i + 1
        
        if 'кафедра' in p.text.lower():
            flag_kaf = True
            vkr_err_point = i + 1 
            
        if 'работа' in p.text.lower() and not flag_vkr:
            flag_vkr = True

            flag_size = True
            flag_bold = True
            
            for run in p.runs:
                if run.font.size != 203200 and flag_size:
                    make_comment(p, "Шрифт должен быть 16 кеггль")
                    flag_size = False
                if not run.bold and flag_bold:
                    make_comment(p, 'Шрифт должен быть полужирным')
                    flag_bold = False
            author_err_point = i + 3
            
        if 'выполн' in p.text.lower():
            flag_author = True
            ruk_err_point = i + 3
            
        if 'руков' in p.text.lower():
            flag_ruk = True

        for city in cities:
            if city.lower() + ' ' in p.text.lower().replace(',', ' ').replace('.', ' '):
                flag_city = True
                break
    
        if re.search(r'(19|[2-9][0-9])\d{2}', p.text):
            flag_year = True

   # if not flag_minest:
    #    make_comment(doc.paragraphs[0], "Не найдена информация о министерстве образования")
    if not flag_un:
        make_comment(doc.paragraphs[un_err_point], "Не найдена информация об учебном заведении")
    if not flag_facul:
        make_comment(doc.paragraphs[facul_err_point], "Не найдена информация о факультете")
    if not flag_kaf:
        make_comment(doc.paragraphs[kaf_err_point], "Не найдена информация о кафедре")
    if not flag_vkr:
        make_comment(doc.paragraphs[vkr_err_point], "Не найдена информация о виде экзаменационной работы")
    if not flag_author:
        make_comment(doc.paragraphs[author_err_point], "Не найдена информация об авторе работы")
    if not flag_ruk:
        make_comment(doc.paragraphs[ruk_err_point], "Не найдена информация о научном руководителе")
    if not flag_city:
        make_comment(doc.paragraphs[i], "Не найдена информация о городе")
    if not flag_year:
        make_comment(doc.paragraphs[i], "Не найдена информация о годе написания работы")

"""Check headings from oglavl block"""
def check_oglavl(doc, page_start_run):
    """Сбор информации из оглавления. Проверка наличия данных элементов в тексте."""
    headers = defaultdict(list)
    if page_start_run[1] is None:
        return headers
    next_page = 2
    while next_page < len(page_start_run) and page_start_run[next_page] is None:
        next_page += 1
    start, end = page_start_run[1][1:], page_start_run[next_page][1:]
    flag_ogl = False
    for i, p in enumerate(tqdm.tqdm(doc.paragraphs[start[0] : end[0]], desc='Process table of contents')):
        if "оглавление" in p.text.lower() or "содержание" in p.text.lower():
            flag_ogl = True
        
        text = p.text.strip()
        if re.search('\d$', text):
            num_page = int(re.search('\d+$', text)[0])
            if re.search('^\d+\.\d+', text):
                headers['level2'].append([p, num_page])
            elif re.search("^\d\.\d\.\d", text):
                headers['level3'].append([p, num_page])
            else:
                headers['level1'].append([p, num_page])

    if not flag_ogl:
        make_comment(doc.paragraphs[start[0]], "Не найдено оглавление работы")
        return
    return headers

def check_headers(doc, headers, page_start_run):
    """Сбор и проверка оформления заголовков разных уровней - основан на информации полученной из оглавления"""
    heading_p = []
    if not headers:
        return

    for headerl1, num_page in headers['level1']:
        start_page_flag = True
        found_heading = False
        if num_page > len(page_start_run) or not page_start_run[num_page - 1]:
            continue
        next_page = num_page
        while next_page < len(page_start_run) and page_start_run[next_page] is None:
            next_page += 1
        right_bound = page_start_run[next_page][1] if next_page < len(page_start_run) else len(doc.paragraphs)
        for i, p in enumerate(doc.paragraphs[page_start_run[num_page - 1][1]: right_bound + 1]):
            p_text = p.text.lower().strip()
            if len(p_text) < 4:
                continue
            elif fuzz.partial_ratio(p_text, headerl1.text.lower()[:len(p_text) + 3]) > 75:
                found_heading = True
                if not start_page_flag:
                    make_comment(p, "Глава должна начинаться с новой страницы")
                if p.style.name != 'Heading 1':
                    make_comment(p, 'Нужен стиль оформления Heading 1')
                heading_p.append(page_start_run[num_page - 1][1] + i)
            else:
                start_page_flag = False
        if not start_page_flag and not found_heading:
            make_comment(headerl1, 'На указанной странице не найдена соответствующая глава')
    
    for headerl2, num_page in headers['level2']:
        found_heading = False
        if num_page > len(page_start_run) or not page_start_run[num_page - 1]:
            continue
        next_page = num_page
        while next_page < len(page_start_run) and page_start_run[next_page] is None:
            next_page += 1
        right_bound = page_start_run[next_page][1] if next_page < len(page_start_run) else len(doc.paragraphs)

        for i, p in enumerate(doc.paragraphs[page_start_run[num_page - 1][1]: right_bound + 1]):
            p_text = p.text.lower().strip()
            if len(p_text) < 4:
                continue
            elif fuzz.partial_ratio(p_text, headerl2.text.lower()[:len(p_text) + 3]) > 75:
                found_heading = True
                if p.style.name != 'Heading 2':
                    make_comment(p, 'Нужен стиль оформления Heading 2')
                heading_p.append(page_start_run[num_page - 1][1] + i)
        if page_start_run[num_page] is not None and not found_heading:
            #print(i, p.text.lower().strip(), headerl2.text.lower()[:len(p_text) + 3], fuzz.partial_ratio(p.text.lower(), headerl2.text.lower()[:len(p.text) + 3]))
            make_comment(headerl2, 'На указанной странице не найдена соответствующая глава')

    for headerl3, num_page in headers['level3']:
        found_heading = False
        if num_page > len(page_start_run) or not page_start_run[num_page - 1]:
            continue
        next_page = num_page
        while next_page < len(page_start_run) and page_start_run[next_page] is None:
            next_page += 1
        right_bound = page_start_run[next_page][1] if next_page < len(page_start_run) else len(doc.paragraphs)

        for i, p in enumerate(doc.paragraphs[page_start_run[num_page - 1][1]: right_bound + 1]):
            p_text = p.text.lower().strip()
            if len(p_text) < 4:
                continue
            elif fuzz.partial_ratio(p_text, headerl3.text.lower()[:len(p_text) + 3]) > 75:
                found_heading = True
                heading_p.append(page_start_run[num_page - 1][1] + i)
        if page_start_run[num_page] is not None and not found_heading:
            make_comment(headerl3, 'На указанной странице не найдена соответствующая глава')
    return heading_p

"""Orfography"""    
def spellchecker(doc):
    """Проверка орфографии текста"""
    checker = SpellChecker(language='ru')
    for p in tqdm.tqdm(doc.paragraphs, desc='Check orfography'):
        try:
            r = requests.get('http://speller.yandex.net/services/spellservice.json/checkText', params={'text': p.text, 'lang': 'ru'})
            if r.status_code == 200:
                if len(r.json()) > 0:
                    out = r.json()[0]
                    word = out['word']
                    variants = [v for v in out['s']]
                    if len(variants):
                        make_comment(p, f'Неизвестное слово {word}. Возможные варинаты: {variants}')
                    else:
                        make_comment(p, f'Неизвестное слово {word}.')
                    #    print(token, variants)
            
            # misspelled = checker.unknown(tokens)
            # for word in misspelled:
                #    print(word, checker.candidates(word))
        except:
            continue

"""Repetition block"""
def normalize_word(morph, word):
    parsed_word = morph.parse(word)[0]
    return parsed_word.normal_form

def check_lexical_repetition(paragraph, morph, threshold=2):
    cleaned_paragraph = re.sub(r'[^\w\s]', '', paragraph.lower())
    
    stop_words = set(stopwords.words('russian'))
    
    words = [normalize_word(morph, word) for word in cleaned_paragraph.split() if word not in stop_words]

    word_counts = Counter(words)

    frequent_words = [word for word, count in word_counts.items() if count >= threshold]

    return frequent_words

def check_ngram_repetition(paragraph, n, morph, threshold=2):
    cleaned_paragraph = re.sub(r'[^\w\s]', '', paragraph.lower())

    stop_words = set(stopwords.words('russian'))
    
    words = [normalize_word(morph, word) for word in cleaned_paragraph.split() if word not in stop_words]

    n_grams = list(ngrams(words, n))

    n_gram_counts = Counter(n_grams)

    frequent_ngrams = [ngram for ngram, count in n_gram_counts.items() if count >= threshold]

    return frequent_ngrams

def check_repetitions(doc):
    """Поиск частых лексических повторов. Вычищается мусор и не учитываются стоп-слова.
      check_ngram_repetition и check_lexical_repetition - вспомогательные функции для проверки повторов слов и нграмм"""
    morph = pymorphy2.MorphAnalyzer()
    for par in tqdm.tqdm(doc.paragraphs, desc='Check repetitions'):
        freq_words = check_lexical_repetition(par.text, morph, threshold=4)
        freq_bigrams = check_ngram_repetition(par.text, 2, morph, threshold=4)
        freq_trigrams = check_ngram_repetition(par.text, 3, morph, threshold=4)
        if freq_trigrams:
            make_comment(par, f'Часто повторяются фразы: {freq_trigrams}')
        elif freq_bigrams:
            make_comment(par, f'Часто повторяются фразы: {freq_bigrams}')
        elif freq_words:
            make_comment(par, f'Часто повторяются слова: {freq_words}')

def get_text_stats(text):
    return BasicStats(text)

def get_readability_stats(text):
    return ReadabilityStats(text)

def get_diversity_stats(text):
    return DiversityStats(text)

def check_readability(ps):
    for p in tqdm.tqdm(ps, desc='Check readability'):
        if len(p.text.strip()) > 30:
            reading_stats = get_readability_stats(p.text).get_stats()
            if reading_stats['flesch_kincaid_grade'] > 18:
                make_comment(p, 'Трудночитаемый текст')
            #elif reading_stats['flesch_reading_easy'] > 80:
            #    make_comment(p, 'Слишком простой текст')

if __name__ == '__main__':
    path = input("PATH: ")
    if not path:
        path = "С.А.+Есенин+Диплом+ВКР.docx"
    if not path[-4:] == 'docx':
        print("Неверный формат файла")
        exit(0)
    doc = Document(path)

    spellchecker(doc)
    splitted_text = split_by_pages(path)
    page_start_run = []
    curr_page = 0
    paragraphs = doc.paragraphs
    blocks = check_main_elements(paragraphs)

    # получение информации о разбиении по страницам и постановка в соответствие данного разбиения и документа в виде обьекта из библиотеки
    for i, p in enumerate(doc.paragraphs):
        while curr_page < len(splitted_text) and splitted_text[curr_page].strip() == '':
            page_start_run.append(None)
            curr_page += 1
        if curr_page == len(splitted_text):
            break
        n_run = 0
        while n_run < len(p.runs):
            if curr_page == len(splitted_text):
                break
            p_text = ''
            start_run = n_run
            while len(p_text.strip()) < 4:
                if n_run == len(p.runs):
                    break
                p_text += p.runs[n_run].text
                if not p_text.replace("\n", " ").replace("\r", " ").replace("\t", " ").strip() and start_run == n_run:
                    start_run += 1
                n_run += 1
            else:
                p_text = p_text.replace("\n", " ").replace("\r", " ").replace("\t", " ").strip()
                if fuzz.partial_ratio(p_text, splitted_text[curr_page][:len(p_text) + 3]) > 80:
                    page_start_run.append((p.runs[start_run], i, start_run))
                    curr_page += 1
                elif curr_page + 1 < len(splitted_text) and fuzz.partial_ratio(p_text, splitted_text[curr_page + 1][:len(p_text) + 3]) > 80:
                    page_start_run.append(None)
                    page_start_run.append((p.runs[start_run], i, start_run))
                    curr_page += 1
                n_run = start_run + 1

    check_title_page(doc, page_start_run)
    """
    with open('pages.txt', 'w', encoding='utf-8') as fout:
        for i, p in enumerate(page_start_run):
            print(i, doc.paragraphs[p[1]].text, file=fout)

    """

    headers = check_oglavl(doc, page_start_run)
    heading_p = check_headers(doc, headers, page_start_run)

    semantic_processing(doc, heading_p)

    check_section(doc)

    start_p = page_start_run[2][1]

    for i, p in tqdm.tqdm(enumerate(doc.paragraphs[start_p:]), desc='Check font and aligment'):
        check_font(p)
        if not heading_p or start_p + i not in heading_p:
            check_aligment(p)

    for i, page_text in tqdm.tqdm(enumerate(splitted_text[2:]), desc='Check page len'):
        if page_text and 0 < len(page_text) < 200:
            try:
                doc.paragraphs[page_start_run[i + 2][1]].add_footnote('Слишком мало информации на странице')
            except:
                pass

    check_repetitions(doc)
    check_readability(doc.paragraphs[start_p:])
    
    """
    for i, p in tqdm.tqdm(enumerate(doc.paragraphs[start_p: -3]), desc='Check sent len'):
        num_short_sent = 0
        for sent in sent_tokenize(p.text):
            if sent and 1 <= len([word for word in word_tokenize(sent) if len(word) > 4]) <= 5:
                num_short_sent += 1
                print(sent)
        if num_short_sent >= 2:
            make_comment(p, 'Слишком много коротких предложений')
    """
    doc.save('outp.docx')
    print("Result in outp.docx")
